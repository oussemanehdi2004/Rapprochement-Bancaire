#!/usr/bin/env python3
"""
Générateur de compte rendu d'avancement du projet BankMatch
Crée un document DOCX professionnel avec captures d'écran et analyse de développement
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from datetime import datetime
import os

class CompteRenduGenerator:
    def __init__(self, output_path="Compte_Rendu_Avancement_BankMatch_Professionnel.docx"):
        self.doc = Document()
        self.output_path = output_path
        self.project_root = r"C:\Users\user\OneDrive\Desktop\rapprochement-bancaire"
        self.setup_styles()
        
    def setup_styles(self):
        """Configure les styles du document"""
        # Style du titre principal
        try:
            title_style = self.doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        except:
            title_style = self.doc.styles['CustomTitle']
        title_font = title_style.font
        title_font.name = 'Calibri'
        title_font.size = Pt(20)
        title_font.bold = True
        title_font.color.rgb = RGBColor(41, 88, 145)  # Bleu professionnel
        
        # Style des titres de section
        try:
            heading_style = self.doc.styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
        except:
            heading_style = self.doc.styles['CustomHeading']
        heading_font = heading_style.font
        heading_font.name = 'Calibri'
        heading_font.size = Pt(16)
        heading_font.bold = True
        heading_font.color.rgb = RGBColor(68, 114, 196)  # Bleu moyen
        
        # Style des sous-titres
        try:
            subheading_style = self.doc.styles.add_style('CustomSubHeading', WD_STYLE_TYPE.PARAGRAPH)
        except:
            subheading_style = self.doc.styles['CustomSubHeading']
        subheading_font = subheading_style.font
        subheading_font.name = 'Calibri'
        subheading_font.size = Pt(13)
        subheading_font.bold = True
        subheading_font.color.rgb = RGBColor(89, 89, 89)  # Gris foncé
        
        # Style du texte normal
        normal_style = self.doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Calibri'
        normal_font.size = Pt(11)
        
        # Style pour les éléments importants
        try:
            highlight_style = self.doc.styles.add_style('Highlight', WD_STYLE_TYPE.PARAGRAPH)
        except:
            highlight_style = self.doc.styles['Highlight']
        highlight_font = highlight_style.font
        highlight_font.name = 'Calibri'
        highlight_font.size = Pt(11)
        highlight_font.bold = True
        highlight_font.color.rgb = RGBColor(192, 0, 0)  # Rouge pour l'emphase
        
    def add_title(self, text):
        """Ajoute un titre principal"""
        paragraph = self.doc.add_paragraph(text)
        paragraph.style = 'CustomTitle'
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph()  # Espace
        
    def add_heading(self, text, level=1):
        """Ajoute un titre de section"""
        paragraph = self.doc.add_paragraph(text)
        if level == 1:
            paragraph.style = 'CustomHeading'
        else:
            paragraph.style = 'CustomSubHeading'
        self.doc.add_paragraph()  # Espace
        
    def add_paragraph(self, text, bold=False, style=None):
        """Ajoute un paragraphe"""
        if style:
            paragraph = self.doc.add_paragraph(text, style=style)
        else:
            paragraph = self.doc.add_paragraph(text)
        if bold:
            paragraph.runs[0].bold = True
        return paragraph
        
    def add_bullet_point(self, text):
        """Ajoute un point de liste"""
        paragraph = self.doc.add_paragraph(text, style='List Bullet')
        return paragraph
        
    def add_image(self, image_path, width=Inches(6.0), caption=None):
        """Ajoute une image avec légende optionnelle"""
        if os.path.exists(image_path):
            try:
                paragraph = self.doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(image_path, width=width)
                
                if caption:
                    caption_para = self.doc.add_paragraph(caption)
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_para.runs[0].italic = True
                    caption_para.runs[0].font.size = Pt(9)
                    
                self.doc.add_paragraph()  # Espace après l'image
                return True
            except Exception as e:
                print(f"Erreur lors de l'ajout de l'image {image_path}: {e}")
                return False
        else:
            print(f"Image non trouvée: {image_path}")
            return False
            
    def add_table(self, headers, data):
        """Ajoute un tableau"""
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # En-têtes
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
            
        # Données
        for row_data in data:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
                
        self.doc.add_paragraph()  # Espace après le tableau
        
    def add_horizontal_line(self):
        """Ajoute une ligne horizontale"""
        paragraph = self.doc.add_paragraph()
        paragraph.add_run("─" * 80)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        self.doc.add_paragraph()
        
    def generate_content(self):
        """Génère le contenu du compte rendu"""
        
        # Page de titre
        self.add_title("COMPTE RENDU D'AVANCEMENT")
        self.add_title("PROJET BANKMATCH - AI & MULTI-BANKING")
        self.add_paragraph("Plateforme de Rapprochement Bancaire Intelligent", bold=True)
        
        date_str = datetime.now().strftime("%d/%m/%Y")
        self.add_paragraph(f"Date: {date_str}", bold=True)
        self.add_paragraph("Auteur: Équipe de Développement BankMatch", bold=True)
        self.add_paragraph("Version: 1.0", bold=True)
        self.add_paragraph("Classification: Confidentiel", bold=True)
        
        self.add_horizontal_line()
        
        # Table des matières simplifiée
        self.add_heading("TABLE DES MATIÈRES")
        self.add_bullet_point("1. Introduction et Contexte du Projet")
        self.add_bullet_point("2. Architecture Technique et Stack Technologique")
        self.add_bullet_point("3. État d'Avancement Détaillé par Module")
        self.add_bullet_point("4. FonctionnalitésImplémentées et Statut")
        self.add_bullet_point("5. Captures d'Écran et Démonstrations")
        self.add_bullet_point("6. Analyse de la Qualité du Code")
        self.add_bullet_point("7. Prochaines Étapes et Roadmap")
        self.add_bullet_point("8. Conclusion et Recommandations")
        
        self.add_horizontal_line()
        
        # Section 1: Introduction
        self.add_heading("1. INTRODUCTION ET CONTEXTE DU PROJET")
        self.add_paragraph(
            "Ce document présente l'état d'avancement du projet BankMatch, "
            "une plateforme innovante de rapprochement bancaire intelligent intégrant "
            "l'intelligence artificielle pour la détection de fraude et le "
            "traitement multi-banking. Le projet s'articule autour de trois "
            "modules principaux conçus selon une architecture microservices."
        )
        self.add_paragraph("Objectifs Principaux:", bold=True)
        self.add_bullet_point("Automatiser le rapprochement bancaire multi-format")
        self.add_bullet_point("Détecter les fraudes transactionnelles par IA")
        self.add_bullet_point("Fournir une interface utilisateur intuitive et responsive")
        self.add_bullet_point("Garantir scalabilité et performance en production")
        
        self.add_paragraph("Modules Principaux:", bold=True)
        self.add_bullet_point("Module Multi-Banking: Parsing et ingestion de fichiers bancaires (CSV, CAMT.053, MT940)")
        self.add_bullet_point("Module Fraud Detection: Analyse IA par Random Forest et explication SHAP")
        self.add_bullet_point("Module Frontend: Interface utilisateur Angular 21 avec Server-Side Rendering")
        
        # Section 2: Architecture
        self.add_heading("2. ARCHITECTURE TECHNIQUE ET STACK TECHNOLOGIQUE")
        
        self.add_paragraph("Architecture Microservices:", bold=True)
        self.add_paragraph(
            "Le projet adopte une architecture microservices avec communication API-First, "
            "assurant isolation, scalabilité et maintenabilité. Chaque service opère "
            "indépendamment avec sa propre base de données et logique métier."
        )
        
        self.add_paragraph("Services Backend:", bold=True)
        backend_data = [
            ["Service", "Technologie", "Port", "Responsabilités"],
            ["Multi-Banking", "FastAPI + Python 3.13", "8010", "Parsing CSV, CAMT.053, MT940, Validation"],
            ["Fraud Detection", "FastAPI + Python 3.13", "8005", "IA Random Forest, SHAP, Règles, Graphe Neo4j"],
            ["Frontend", "Angular 21 + SSR", "4200", "Interface utilisateur, Dashboard, Visualisation"]
        ]
        self.add_table(["Service", "Technologie", "Port", "Responsabilités"], backend_data)
        
        # Dépendances principales
        self.add_paragraph("Dépendances Principales:", bold=True)
        deps_data = [
            ["Module", "Dépendances Clés", "Cas d'Usage"],
            ["Multi-Banking", "fastapi, uvicorn, lxml, httpx, pytest", "API haute performance, parsing XML"],
            ["Fraud Detection", "scikit-learn, shap, neo4j, supabase", "Machine Learning, Graph DB, Persistance"],
            ["Frontend", "Angular 21, TailwindCSS, Axios", "UI moderne, HTTP client, Styling"]
        ]
        self.add_table(["Module", "Dépendances Clés", "Cas d'Usage"], deps_data)
        
        self.add_paragraph("Pattern d'Authentification:", bold=True)
        self.add_paragraph(
            "Architecture de sécurité avec JWT pour service-to-service et utilisateur-final. "
            "Tokens internes avec durée de vie courte (30s) pour communication inter-services."
        )
        
        # Section 3: État d'avancement
        self.add_heading("3. ÉTAT D'AVANCEMENT DÉTAILLÉ PAR MODULE")
        
        # Module Multi-Banking
        self.add_paragraph("Module Multi-Banking - État: 85% Complet", bold=True)
        self.add_bullet_point("✅ Parsing CSV implémenté et testé unitairement")
        self.add_bullet_point("✅ Parsing CAMT.053 (ISO 20022) implémenté et testé")
        self.add_bullet_point("✅ Parsing MT940 (SWIFT) implémenté et testé")
        self.add_bullet_point("✅ Validation des transactions avec règles métier")
        self.add_bullet_point("✅ Authentification service-to-service JWT avec secret dédié")
        self.add_bullet_point("✅ Intégration Fraud Detection (endpoint /ingest)")
        self.add_bullet_point("✅ Logging structuré JSON avec request ID tracking")
        self.add_bullet_point("✅ Health check endpoint pour monitoring")
        self.add_bullet_point("✅ Configuration Docker avec HEALTHCHECK")
        self.add_bullet_point("⏳ Intégration BankMatch API (en attente de finalisation contrat)")
        
        # Module Fraud Detection
        self.add_paragraph("Module Fraud Detection - État: 90% Complet", bold=True)
        self.add_bullet_point("✅ Modèle Random Forest entraîné sur données bancaires")
        self.add_bullet_point("✅ Explicabilité SHAP intégrée pour transparence IA")
        self.add_bullet_point("✅ Moteur de règles métier configurable dynamiquement")
        self.add_bullet_point("✅ Analyse de graphe Neo4j pour détection de réseaux frauduleux")
        self.add_bullet_point("✅ Persistance Supabase avec schema transaction_reference")
        self.add_bullet_point("✅ Authentification interne JWT avec validation")
        self.add_bullet_point("✅ Logging structuré avec traçabilité cross-service")
        self.add_bullet_point("✅ Correction champ transaction_reference (ex: mongo_transaction_id)")
        self.add_bullet_point("✅ CORS configuré dynamiquement par environnement")
        self.add_bullet_point("✅ API endpoints /analyze, /rules, /graph configurés")
        self.add_bullet_point("⏳ Intégration backend BankMatch (en attente)")
        
        # Module Frontend
        self.add_paragraph("Module Frontend - État: 70% Complet", bold=True)
        self.add_bullet_point("✅ Architecture Angular 21 avec Server-Side Rendering")
        self.add_bullet_point("✅ Interface de détection de fraude avec dashboard")
        self.add_bullet_point("✅ Composants header et sidebar navigation")
        self.add_bullet_point("✅ Services API générés automatiquement (OpenAPI)")
        self.add_bullet_point("✅ Configuration TailwindCSS pour styling responsive")
        self.add_bullet_point("✅ Auth interceptor pour gestion tokens JWT")
        self.add_bullet_point("✅ Structure modulaire par feature (fraud-detection, transactions)")
        self.add_bullet_point("⏳ Intégration complète avec backend BankMatch")
        self.add_bullet_point("⏳ Tests E2E avec Playwright ou Cypress")
        
        # Section 4: Fonctionnalités
        self.add_heading("4. FONCTIONNALITÉS IMPLÉMENTÉES ET STATUT")
        
        features_data = [
            ["Fonctionnalité", "Statut", "Module", "Priorité"],
            ["Parsing multi-format (CSV, CAMT.053, MT940)", "Complet", "Multi-Banking", "Critique"],
            ["Validation transactions avec règles", "Complet", "Multi-Banking", "Critique"],
            ["Détection de fraude IA Random Forest", "Complet", "Fraud Detection", "Critique"],
            ["Explicabilité SHAP (transparence IA)", "Complet", "Fraud Detection", "Élevée"],
            ["Moteur de règles configurable", "Complet", "Fraud Detection", "Élevée"],
            ["Analyse de graphe Neo4j", "Complet", "Fraud Detection", "Moyenne"],
            ["Authentification JWT service-to-service", "Complet", "Tous", "Critique"],
            ["Logging structuré JSON", "Complet", "Tous", "Élevée"],
            ["Persistance Supabase", "Complet", "Fraud Detection", "Critique"],
            ["Interface utilisateur responsive", "En cours", "Frontend", "Élevée"],
            ["Dashboard visualisation temps réel", "En cours", "Frontend", "Moyenne"],
            ["Tests E2E pipeline complet", "À faire", "Tous", "Moyenne"]
        ]
        self.add_table(["Fonctionnalité", "Statut", "Module", "Priorité"], features_data)
        
        # Section 5: Captures d'écran
        self.add_heading("5. CAPTURES D'ÉCRAN ET DÉMONSTRATIONS")
        
        self.add_paragraph(
            "Cette section présente les interfaces utilisateur et démonstrations "
            "des fonctionnalités implémentées. Les captures illustrent le flux "
            "utilisateur complet depuis l'ingestion des fichiers jusqu'à l'analyse "
            "de fraude et la visualisation des résultats."
        )
        
        # Captures fraud detection
        self.add_paragraph("Interface de Détection de Fraude - Dashboard Principal:", bold=True)
        
        fraud_captures = [
            r"C:\Users\user\OneDrive\Desktop\rapprochement-bancaire\fraud-detection\assets\captures\image (1).png",
            r"C:\Users\user\OneDrive\Desktop\rapprochement-bancaire\fraud-detection\assets\captures\image (2).png",
            r"C:\Users\user\OneDrive\Desktop\rapprochement-bancaire\fraud-detection\assets\captures\image (3).png",
            r"C:\Users\user\OneDrive\Desktop\rapprochement-bancaire\fraud-detection\assets\captures\visualisation.png"
        ]
        
        captions = [
            "Figure 1: Dashboard principal avec alertes de fraude en temps réel",
            "Figure 2: Détail d'une transaction suspecte avec analyse SHAP",
            "Figure 3: Interface de configuration des règles métier",
            "Figure 4: Visualisation graphique des réseaux de transactions"
        ]
        
        for i, (capture_path, caption) in enumerate(zip(fraud_captures, captions)):
            if self.add_image(capture_path, width=Inches(5.5), caption=caption):
                pass
        
        # Captures multi-banking
        self.add_paragraph("Interface Multi-Banking - Ingestion et Parsing:", bold=True)
        
        banking_captures = [
            r"C:\Users\user\OneDrive\Desktop\rapprochement-bancaire\multi-banking\captures\30ace81a-68fa-458f-bb20-49b86803e174.png",
            r"C:\Users\user\OneDrive\Desktop\rapprochement-bancaire\multi-banking\captures\3f5cc2a3-86a1-4e39-ad1f-9e99184ef9b1.png",
            r"C:\Users\user\OneDrive\Desktop\rapprochement-bancaire\multi-banking\captures\ae9c0ef9-5e47-4bf3-86bc-9758daf0350d.png"
        ]
        
        banking_captions = [
            "Figure 5: Interface d'upload de fichiers bancaires multi-format",
            "Figure 6: Résultats du parsing avec validation des transactions",
            "Figure 7: Intégration avec le service de détection de fraude"
        ]
        
        for i, (capture_path, caption) in enumerate(zip(banking_captures, banking_captions)):
            if self.add_image(capture_path, width=Inches(5.5), caption=caption):
                pass
        
        # Section 6: Analyse du code
        self.add_heading("6. ANALYSE DE LA QUALITÉ DU CODE")
        
        self.add_paragraph("Qualité du Code et Best Practices:", bold=True)
        self.add_bullet_point("✅ Structure modulaire avec séparation claire des responsabilités")
        self.add_bullet_point("✅ Tests unitaires complets pour parseurs (CSV, CAMT.053, MT940)")
        self.add_bullet_point("✅ Tests d'intégration avec mocks pour API externes")
        self.add_bullet_point("✅ Logging structuré JSON pour traçabilité cross-service")
        self.add_bullet_point("✅ Gestion d'erreurs robuste avec HTTPException personnalisées")
        self.add_bullet_point("✅ Configuration par variables d'environnement (.env)")
        self.add_bullet_point("✅ Documentation inline type hints et docstrings")
        self.add_bullet_point("✅ Code style PEP 8 respecté")
        self.add_bullet_point("✅ Schémas Pydantic pour validation des données")
        
        self.add_paragraph("Principes Architecturaux Appliqués:", bold=True)
        self.add_bullet_point("✅ Communication API-First entre microservices")
        self.add_bullet_point("✅ Isolation des services avec secrets JWT dédiés")
        self.add_bullet_point("✅ Authentification JWT avec tokens internes à durée courte")
        self.add_bullet_point("✅ Pas d'accès direct MongoDB - communication via API BankMatch")
        self.add_bullet_point("✅ CORS configuré dynamiquement par environnement")
        self.add_bullet_point("✅ Request ID tracking pour traçabilité distribuée")
        self.add_bullet_point("✅ Health checks pour monitoring orchestration")
        
        self.add_paragraph("Métriques de Qualité:", bold=True)
        quality_data = [
            ["Métrique", "Status", "Détails"],
            ["Couverture Tests", "En cours", "Tests unitaires parseurs, tests intégration API"],
            ["Documentation", "Bonne", "Docstrings, OpenAPI/Swagger disponibles"],
            ["Complexité Cyclomatique", "Faible", "Fonctions courtes et ciblées"],
            ["Dépendances", "Stables", "Versions figées dans requirements.txt"],
            ["Sécurité", "Renforcée", "JWT, CORS, validation entrées"]
        ]
        self.add_table(["Métrique", "Status", "Détails"], quality_data)
        
        # Section 7: Prochaines étapes
        self.add_heading("7. PROCHAINES ÉTAPES ET ROADMAP")
        
        self.add_paragraph("Intégration BankMatch - Priorité Critique:", bold=True)
        self.add_bullet_point("⏳ Finaliser contrat API /api/import avec équipe BankMatch")
        self.add_bullet_point("⏳ Finaliser contrat API /reconciliation/sessions/:id/matching/start")
        self.add_bullet_point("⏳ Implémenter validation tokens internes BankMatch côté backend")
        self.add_bullet_point("⏳ Activer BANKMATCH_INTEGRATION_ENABLED=true en production")
        self.add_bullet_point("⏳ Valider flux complet: Multi-Banking → BankMatch → Fraud Detection")
        
        self.add_paragraph("Déploiement et Infrastructure - Priorité Élevée:", bold=True)
        self.add_bullet_point("⏳ Configurer limites ressources Docker (CPU, memory)")
        self.add_bullet_point("⏳ Implémenter CI/CD pipeline (GitHub Actions ou GitLab CI)")
        self.add_bullet_point("⏳ Configurer monitoring Prometheus + Grafana")
        self.add_bullet_point("⏳ Configurer traçage distribué Jaeger ou Zipkin")
        self.add_bullet_point("⏳ Configurer secrets manager (HashiCorp Vault ou AWS Secrets Manager)")
        self.add_bullet_point("⏳ Configurer mutual TLS pour communication service-to-service")
        
        self.add_paragraph("Frontend et UX - Priorité Élevée:", bold=True)
        self.add_bullet_point("⏳ Finaliser intégration backend BankMatch")
        self.add_bullet_point("⏳ Implémenter tests E2E avec Playwright ou Cypress")
        self.add_bullet_point("⏳ Optimisation performance SSR (Server-Side Rendering)")
        self.add_bullet_point("⏳ Accessibilité WCAG 2.1 compliance")
        self.add_bullet_point("⏳ Internationalisation (i18n) pour multi-langue")
        
        self.add_paragraph("Tests et Qualité - Priorité Moyenne:", bold=True)
        self.add_bullet_point("⏳ Atteindre 80% couverture de tests unitaires")
        self.add_bullet_point("⏳ Tests de charge et performance (Locust ou k6)")
        self.add_bullet_point("⏳ Tests de sécurité (SAST/DAST avec SonarQube)")
        self.add_bullet_point("⏳ Tests de pénétration pour API endpoints")
        
        # Section 8: Conclusion
        self.add_heading("8. CONCLUSION ET RECOMMANDATIONS")
        
        self.add_paragraph("Synthèse de l'État d'Avancement:", bold=True)
        self.add_paragraph(
            "Le projet BankMatch a atteint un stade avancé de développement avec "
            "l'implémentation complète des modules core (Multi-Banking et Fraud Detection). "
            "L'architecture microservices est solide, respecte les best pratiques modernes, "
            "et est prête pour l'intégration finale avec le backend BankMatch. "
            "Les fondations techniques sont robustes et permettent une évolution scalable."
        )
        
        self.add_paragraph("Points Forts Identifiés:", bold=True)
        self.add_bullet_point("✅ Architecture microservices bien conçue et isolée")
        self.add_bullet_point("✅ Implémentation IA avec explicabilité (SHAP)")
        self.add_bullet_point("✅ Logging structuré pour observabilité")
        self.add_bullet_point("✅ Authentification sécurisée JWT")
        self.add_bullet_point("✅ Tests unitaires couvrant les composants critiques")
        self.add_bullet_point("✅ Configuration flexible par environnement")
        
        self.add_paragraph("Recommandations Prioritaires:", bold=True)
        self.add_bullet_point("🎯 Finaliser dans les meilleurs délais l'intégration API BankMatch")
        self.add_bullet_point("🎯 Mettre en place pipeline CI/CD pour déploiement continu")
        self.add_bullet_point("🎯 Implémenter monitoring complet avant mise en production")
        self.add_bullet_point("🎯 Renforcer couverture de tests (E2E, performance, sécurité)")
        self.add_bullet_point("🎯 Documenter procédures opérationnelles et runbooks")
        
        self.add_paragraph("Risques et Atténuations:", bold=True)
        risks_data = [
            ["Risque", "Impact", "Atténuation"],
            ["Dépendance API BankMatch", "Critique", "Documentation contrat détaillé, mocks robustes"],
            ["Performance scaling", "Élevé", "Tests charge, cache Redis, database optimization"],
            ["Sécurité données", "Critique", "Audit sécurité, encryption at rest, TLS mutual"],
            ["Complexité ops", "Moyen", "Infrastructure as Code, monitoring proactif"]
        ]
        self.add_table(["Risque", "Impact", "Atténuation"], risks_data)
        
        # Conclusion finale
        self.add_horizontal_line()
        self.add_paragraph(
            "En conclusion, le projet BankMatch présente une base technique solide "
            "et innovante pour le rapprochement bancaire intelligent. La qualité du code, "
            "l'architecture modernes et l'approche IA positionnent favorablement le projet "
            "pour une adoption réussie en production. Les prochaines semaines seront "
            "critiques pour finaliser l'intégration et préparer le déploiement."
        )
        
        # Signature
        self.doc.add_paragraph()
        signature_para = self.doc.add_paragraph()
        signature_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        signature_para.add_run("_____________________")
        signature_para.add_run("\nResponsable Technique")
        signature_para.add_run("\nÉquipe BankMatch")
        signature_para.add_run(f"\nDate: {datetime.now().strftime('%d/%m/%Y')}")
        
    def save(self):
        """Sauvegarde le document"""
        try:
            self.doc.save(self.output_path)
            print(f"Document généré avec succès: {self.output_path}")
            return True
        except Exception as e:
            print(f"Erreur lors de la sauvegarde: {e}")
            return False

def main():
    generator = CompteRenduGenerator()
    generator.generate_content()
    generator.save()

if __name__ == "__main__":
    main()